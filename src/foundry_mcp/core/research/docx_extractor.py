"""DOCX text extraction for deep research workflows.

Provides secure DOCX text extraction with paragraph and table content
extraction. Uses python-docx as the extraction engine.

Security Features:
    - SSRF protection: Blocks internal IPs, localhost, and private networks
    - Magic byte validation: Verifies PK (ZIP) header + word/document.xml entry
    - Content-type validation: Checks HTTP response content-type
    - Size limits: Configurable maximum DOCX size

Key Components:
    - DocxExtractionResult: Dataclass containing extracted text and metadata
    - DocxExtractor: Main class for extracting text from DOCX files/bytes

Usage:
    from foundry_mcp.core.research.docx_extractor import (
        DocxExtractor,
        DocxExtractionResult,
    )

    # Create extractor
    extractor = DocxExtractor()

    # Extract from bytes
    result = await extractor.extract(docx_bytes)

    # Extract from URL (with SSRF protection)
    result = await extractor.extract_from_url("https://example.com/doc.docx")

    # Access results
    print(result.text)
    print(result.paragraph_count)
    print(result.table_count)
"""

from __future__ import annotations

import asyncio
import io
import logging
import time
import zipfile
from dataclasses import dataclass, field
from typing import Any, Optional, Union
from urllib.parse import urljoin

logger = logging.getLogger(__name__)

# =============================================================================
# Metrics (Optional - graceful degradation if prometheus_client not installed)
# =============================================================================

try:
    from prometheus_client import Counter, Histogram

    _PROMETHEUS_AVAILABLE = True
except ImportError:
    _PROMETHEUS_AVAILABLE = False
    Counter: Any = None
    Histogram: Any = None

# Metrics instances (lazily initialized)
_docx_extraction_duration: Optional[Any] = None
_docx_extraction_paragraphs: Optional[Any] = None
_metrics_initialized: bool = False


def _init_metrics() -> None:
    """Initialize DOCX extraction metrics (thread-safe, idempotent)."""
    global _docx_extraction_duration, _docx_extraction_paragraphs, _metrics_initialized

    if _metrics_initialized or not _PROMETHEUS_AVAILABLE:
        return

    _metrics_initialized = True

    _docx_extraction_duration = Histogram(
        "foundry_mcp_docx_extraction_duration_seconds",
        "DOCX extraction duration in seconds",
        ["status"],
        buckets=(0.1, 0.25, 0.5, 1.0, 2.5, 5.0, 10.0, 30.0, 60.0),
    )

    _docx_extraction_paragraphs = Counter(
        "foundry_mcp_docx_extraction_paragraphs_total",
        "Total number of paragraphs extracted from DOCX files",
        ["status"],
    )

    logger.debug("DOCX extraction metrics initialized")


def _record_extraction_metrics(
    duration_seconds: float,
    paragraphs_extracted: int,
    status: str,
) -> None:
    """Record DOCX extraction metrics.

    Args:
        duration_seconds: Extraction duration in seconds.
        paragraphs_extracted: Number of paragraphs successfully extracted.
        status: Extraction status - "success", "partial", or "failure".
    """
    if not _PROMETHEUS_AVAILABLE:
        return

    _init_metrics()

    if _docx_extraction_duration is not None:
        _docx_extraction_duration.labels(status=status).observe(duration_seconds)

    if _docx_extraction_paragraphs is not None and paragraphs_extracted > 0:
        _docx_extraction_paragraphs.labels(status=status).inc(paragraphs_extracted)


# =============================================================================
# Lazy Import for python-docx
# =============================================================================

_docx_module: Optional[object] = None
_docx_checked: bool = False


def _get_docx_module():
    """Lazy import for python-docx.

    Returns the docx module if available, None otherwise.
    The import is cached after first call to avoid repeated import attempts.

    Returns:
        docx module or None if not installed.
    """
    global _docx_module, _docx_checked

    if _docx_checked:
        return _docx_module

    _docx_checked = True
    try:
        import docx

        _docx_module = docx
        logger.debug("python-docx available for DOCX extraction")
    except ImportError:
        _docx_module = None
        logger.debug("python-docx not installed, DOCX extraction unavailable")

    return _docx_module


# =============================================================================
# Security Constants
# =============================================================================

DOCX_MAGIC_BYTES = b"PK\x03\x04"
"""DOCX files (OOXML ZIP archives) start with this magic byte sequence."""

VALID_DOCX_CONTENT_TYPES = frozenset(
    [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/octet-stream",  # Some servers serve DOCX with this
    ]
)
"""Content-types that are acceptable for DOCX responses."""

DEFAULT_MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10 MB
"""Default maximum DOCX file size in bytes."""

DEFAULT_FETCH_TIMEOUT = 30.0
"""Default timeout for URL fetches in seconds."""

MAX_DOCX_REDIRECTS = 5
"""Maximum number of redirects to follow when fetching DOCX files."""


# =============================================================================
# Error classes (canonical definitions in foundry_mcp.core.errors.research)
# =============================================================================

from foundry_mcp.core.errors.research import (  # noqa: E402
    DocxSecurityError,
    DocxSizeError,
    InvalidDocxError,
)

# Reuse SSRF validation from pdf_extractor
from foundry_mcp.core.research.pdf_extractor import (  # noqa: E402
    SSRFError,
    validate_url_for_ssrf,
)

# =============================================================================
# Validation Functions
# =============================================================================


def validate_docx_magic_bytes(data: bytes) -> None:
    """Validate DOCX magic bytes and structure.

    Checks for ZIP header (PK\\x03\\x04) and probes for the
    ``word/document.xml`` entry that identifies an OOXML Word document.

    Args:
        data: DOCX file data.

    Raises:
        InvalidDocxError: If magic bytes don't match or word/ entries not found.
    """
    if len(data) < len(DOCX_MAGIC_BYTES):
        raise InvalidDocxError(f"Data too short to be a DOCX ({len(data)} bytes)")

    if not data.startswith(DOCX_MAGIC_BYTES):
        preview = data[:20].hex()
        raise InvalidDocxError(f"Invalid DOCX: missing PK header. Got: {preview}...")

    # Probe ZIP for word/ entries to confirm it's a DOCX (not just any ZIP)
    try:
        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            names = zf.namelist()
            if not any(name.startswith("word/") for name in names):
                raise InvalidDocxError(
                    "ZIP archive does not contain word/ entries — not a valid DOCX file"
                )
    except zipfile.BadZipFile as e:
        raise InvalidDocxError(f"Invalid DOCX: corrupt ZIP archive: {e}") from e


def validate_docx_content_type(content_type: str | None) -> None:
    """Validate HTTP content-type for DOCX responses.

    Args:
        content_type: Content-Type header value.

    Raises:
        InvalidDocxError: If content-type is not acceptable for DOCX.
    """
    if not content_type:
        logger.warning("No Content-Type header, proceeding with magic byte validation")
        return

    # Extract base content type (ignore parameters like charset)
    base_type = content_type.split(";")[0].strip().lower()

    if base_type not in VALID_DOCX_CONTENT_TYPES:
        raise InvalidDocxError(
            f"Invalid Content-Type for DOCX: {content_type}. "
            f"Expected one of: {', '.join(VALID_DOCX_CONTENT_TYPES)}"
        )


# =============================================================================
# Result Dataclass
# =============================================================================


@dataclass
class DocxExtractionResult:
    """Result of DOCX text extraction.

    Contains the extracted text from paragraphs and tables, along with
    metadata and any warnings encountered during extraction.

    Attributes:
        text: Concatenated text from all paragraphs and tables.
        warnings: List of warning messages from extraction.
        paragraph_count: Number of paragraphs in the document.
        table_count: Number of tables in the document.
    """

    text: str
    warnings: list[str] = field(default_factory=list)
    paragraph_count: int = 0
    table_count: int = 0

    @property
    def success(self) -> bool:
        """Check if extraction produced any text."""
        return len(self.text.strip()) > 0

    @property
    def has_warnings(self) -> bool:
        """Check if extraction produced any warnings."""
        return len(self.warnings) > 0


# =============================================================================
# Extractor Class
# =============================================================================


class DocxExtractor:
    """Extracts text from DOCX files with security hardening.

    Uses python-docx for text extraction, extracting both paragraph text
    and table cell content.

    Security Features:
        - SSRF protection for URL fetching (blocks internal IPs/localhost)
        - Magic byte + ZIP structure validation
        - Content-type validation for HTTP responses
        - Configurable size limits

    The extractor is designed for async usage in research workflows,
    running CPU-bound python-docx operations in a thread pool to avoid
    blocking the event loop.

    Attributes:
        max_size: Maximum DOCX file size in bytes (default: 10MB).
        timeout: Timeout for URL fetches in seconds (default: 30s).

    Example:
        extractor = DocxExtractor()

        # Extract from bytes (validates magic bytes)
        result = await extractor.extract(docx_bytes)

        # Extract from URL (with SSRF protection)
        result = await extractor.extract_from_url("https://example.com/doc.docx")

        # Extract with custom limits
        extractor = DocxExtractor(max_size=5*1024*1024)
    """

    def __init__(
        self,
        max_size: int = DEFAULT_MAX_DOCX_SIZE,
        fetch_timeout: float = DEFAULT_FETCH_TIMEOUT,
    ):
        """Initialize DocxExtractor with resource limits.

        Args:
            max_size: Maximum DOCX file size in bytes (default: 10MB).
            fetch_timeout: Timeout for URL fetches in seconds (default: 30s).
        """
        self.max_size = max_size
        self.timeout = fetch_timeout

    async def extract(
        self,
        source: Union[bytes, io.BytesIO],
        *,
        validate_magic: bool = True,
    ) -> DocxExtractionResult:
        """Extract text from a DOCX source.

        Validates DOCX structure before parsing and runs extraction in a
        thread pool to avoid blocking the event loop.

        Args:
            source: DOCX content as bytes or BytesIO stream.
            validate_magic: Whether to validate magic bytes and ZIP structure
                (default: True).

        Returns:
            DocxExtractionResult with extracted text and metadata.

        Raises:
            ValueError: If source is not bytes or BytesIO.
            InvalidDocxError: If magic byte or structure validation fails.
            DocxSizeError: If DOCX exceeds max_size.
            RuntimeError: If python-docx is not installed.
        """
        if isinstance(source, bytes):
            docx_bytes = source
        elif isinstance(source, io.BytesIO):
            docx_bytes = source.getvalue()
        else:
            raise ValueError(f"source must be bytes or BytesIO, got {type(source).__name__}")

        # Check size limit
        if len(docx_bytes) > self.max_size:
            raise DocxSizeError(
                f"DOCX size ({len(docx_bytes)} bytes) exceeds limit ({self.max_size} bytes)"
            )

        # Validate magic bytes and ZIP structure
        if validate_magic:
            validate_docx_magic_bytes(docx_bytes)

        # Check python-docx availability
        docx_mod = _get_docx_module()
        if docx_mod is None:
            raise RuntimeError(
                "python-docx is required for DOCX extraction. "
                "Install with: pip install 'foundry-mcp[docx]'"
            )

        # Run CPU-bound extraction in thread pool
        try:
            return await asyncio.wait_for(
                asyncio.to_thread(self._extract_sync, docx_bytes),
                timeout=self.timeout,
            )
        except asyncio.TimeoutError as e:
            raise DocxSecurityError(
                f"DOCX extraction timed out after {self.timeout}s"
            ) from e

    def _extract_sync(self, data: bytes) -> DocxExtractionResult:
        """Synchronous extraction implementation.

        Extracts text from paragraphs and tables using python-docx.

        Args:
            data: Raw DOCX bytes.

        Returns:
            DocxExtractionResult with extracted content.
        """
        start_time = time.perf_counter()
        warnings: list[str] = []

        docx_mod = _get_docx_module()
        assert docx_mod is not None  # Checked in extract()

        try:
            doc = docx_mod.Document(io.BytesIO(data))  # type: ignore[union-attr]
        except Exception as e:
            logger.warning(f"Failed to read DOCX: {e}")
            duration = time.perf_counter() - start_time
            _record_extraction_metrics(duration, 0, "failure")
            return DocxExtractionResult(
                text="",
                warnings=[f"Failed to parse DOCX: {e}"],
                paragraph_count=0,
                table_count=0,
            )

        # Extract paragraph text
        paragraphs = doc.paragraphs  # type: ignore[union-attr]
        paragraph_texts: list[str] = []
        for para in paragraphs:
            text = para.text.strip()
            if text:
                paragraph_texts.append(text)

        paragraph_count = len(paragraphs)

        if not paragraph_texts:
            warnings.append("No text content found in paragraphs")

        # Extract table text
        tables = doc.tables  # type: ignore[union-attr]
        table_texts: list[str] = []
        for table in tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    table_texts.append(" | ".join(row_cells))

        table_count = len(tables)

        # Combine paragraph and table text
        all_parts: list[str] = []
        if paragraph_texts:
            all_parts.append("\n\n".join(paragraph_texts))
        if table_texts:
            all_parts.append("\n".join(table_texts))

        full_text = "\n\n".join(all_parts)

        duration = time.perf_counter() - start_time
        extracted_paragraphs = len(paragraph_texts)

        logger.debug(
            f"Extracted {extracted_paragraphs} paragraphs, "
            f"{table_count} tables, {len(full_text)} chars"
        )

        status = "success" if full_text.strip() else "failure"
        _record_extraction_metrics(duration, extracted_paragraphs, status)

        return DocxExtractionResult(
            text=full_text,
            warnings=warnings,
            paragraph_count=paragraph_count,
            table_count=table_count,
        )

    async def extract_from_url(self, url: str) -> DocxExtractionResult:
        """Extract text from a DOCX at a URL with SSRF protection.

        Validates the URL against SSRF attacks before fetching, then
        validates content-type and magic bytes before extraction.

        Security features:
            - SSRF validation on initial URL
            - SSRF re-validation after redirects
            - Streaming download with early abort at size limit
            - Content-type and magic byte validation

        Args:
            url: URL to fetch the DOCX from. Must be http or https.

        Returns:
            DocxExtractionResult with extracted text and metadata.

        Raises:
            SSRFError: If URL fails SSRF validation.
            InvalidDocxError: If content-type or magic bytes are invalid.
            DocxSizeError: If DOCX exceeds max_size.
        """
        # Validate initial URL for SSRF before any network request
        validate_url_for_ssrf(url)

        try:
            import httpx
        except ImportError as e:
            raise ImportError(
                "httpx is required for URL fetching. Install with: pip install httpx"
            ) from e

        logger.debug(f"Fetching DOCX from URL: {url}")

        current_url = url
        visited: set[str] = set()

        async with httpx.AsyncClient(timeout=self.timeout) as client:
            for _redirect_index in range(MAX_DOCX_REDIRECTS + 1):
                if current_url in visited:
                    raise SSRFError(f"Redirect loop detected for {current_url}")
                visited.add(current_url)

                # Validate URL for SSRF before any network request
                validate_url_for_ssrf(current_url)

                async with client.stream(
                    "GET",
                    current_url,
                    follow_redirects=False,
                    headers={"User-Agent": "foundry-mcp/1.0 DocxExtractor"},
                ) as response:
                    if response.status_code in {301, 302, 303, 307, 308}:
                        location = response.headers.get("location")
                        if not location:
                            raise InvalidDocxError(
                                f"Redirect response missing Location header: {current_url}"
                            )
                        next_url = urljoin(current_url, location)
                        logger.debug("Redirect detected: %s -> %s", current_url, next_url)
                        current_url = next_url
                        continue

                    response.raise_for_status()

                    # Validate content-type
                    content_type = response.headers.get("content-type")
                    validate_docx_content_type(content_type)

                    # Stream content with size limit enforcement
                    chunks: list[bytes] = []
                    total_size = 0

                    async for chunk in response.aiter_bytes(chunk_size=65536):
                        total_size += len(chunk)
                        if total_size > self.max_size:
                            raise DocxSizeError(
                                f"DOCX size exceeds limit ({self.max_size} bytes), "
                                f"download aborted at {total_size} bytes"
                            )
                        chunks.append(chunk)

                    docx_bytes = b"".join(chunks)

                # Validate magic bytes and ZIP structure
                validate_docx_magic_bytes(docx_bytes)

                logger.debug(f"Downloaded {len(docx_bytes)} bytes from {current_url}")

                # Extract text (magic already validated)
                return await self.extract(docx_bytes, validate_magic=False)

        raise InvalidDocxError(
            f"Too many redirects while fetching DOCX (max {MAX_DOCX_REDIRECTS})"
        )
