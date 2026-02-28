"""Tests for DOCX extractor module.

Tests cover:
1. Valid DOCX extraction - successful extraction from bytes
2. Table extraction - extracting table cell text
3. Magic byte validation - rejecting invalid DOCX headers
4. ZIP structure validation - rejecting non-DOCX ZIP files
5. Size limits enforcement - enforcing configurable max size
6. SSRF protection - blocking internal IPs, localhost, private networks
7. Content-type validation - rejecting invalid content types
8. Empty and corrupted document handling
9. Graceful degradation when python-docx is not installed
10. Thread pool execution for CPU-bound work
"""

import io
import zipfile
from unittest.mock import patch

import pytest

from foundry_mcp.core.errors.research import (
    DocxSecurityError,
    DocxSizeError,
    InvalidDocxError,
)
from foundry_mcp.core.research.docx_extractor import (
    DEFAULT_MAX_DOCX_SIZE,
    DOCX_MAGIC_BYTES,
    DocxExtractionResult,
    DocxExtractor,
    validate_docx_content_type,
    validate_docx_magic_bytes,
)
from foundry_mcp.core.research.pdf_extractor import SSRFError

# =============================================================================
# Fixtures
# =============================================================================


def _create_docx_bytes(paragraphs: list[str] | None = None, tables: list[list[list[str]]] | None = None) -> bytes:
    """Create a DOCX file in memory using python-docx.

    Args:
        paragraphs: List of paragraph text strings.
        tables: List of tables, each table is a list of rows, each row is a list of cell strings.

    Returns:
        DOCX file as bytes.
    """
    import docx

    doc = docx.Document()

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            if not table_data:
                continue
            cols = max(len(row) for row in table_data)
            table = doc.add_table(rows=len(table_data), cols=cols)
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def simple_docx_bytes() -> bytes:
    """Create minimal DOCX with paragraph text."""
    return _create_docx_bytes(paragraphs=["Hello World", "This is a test document."])


@pytest.fixture
def docx_with_tables_bytes() -> bytes:
    """Create DOCX with both paragraphs and tables."""
    return _create_docx_bytes(
        paragraphs=["Document with tables"],
        tables=[
            [
                ["Name", "Age", "City"],
                ["Alice", "30", "New York"],
                ["Bob", "25", "San Francisco"],
            ]
        ],
    )


@pytest.fixture
def empty_docx_bytes() -> bytes:
    """Create a DOCX with no content."""
    return _create_docx_bytes()


@pytest.fixture
def generic_zip_bytes() -> bytes:
    """Create a generic ZIP file (not a DOCX)."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("data.txt", "some data")
    return buf.getvalue()


# =============================================================================
# DocxExtractionResult Tests
# =============================================================================


class TestDocxExtractionResult:
    """Tests for the DocxExtractionResult dataclass."""

    def test_success_with_text(self):
        result = DocxExtractionResult(text="Hello", paragraph_count=1)
        assert result.success is True

    def test_success_without_text(self):
        result = DocxExtractionResult(text="", paragraph_count=0)
        assert result.success is False

    def test_success_whitespace_only(self):
        result = DocxExtractionResult(text="   \n  ", paragraph_count=1)
        assert result.success is False

    def test_has_warnings(self):
        result = DocxExtractionResult(text="", warnings=["Something happened"])
        assert result.has_warnings is True

    def test_no_warnings(self):
        result = DocxExtractionResult(text="Hello")
        assert result.has_warnings is False


# =============================================================================
# Magic Byte Validation Tests
# =============================================================================


class TestValidateDocxMagicBytes:
    """Tests for validate_docx_magic_bytes()."""

    def test_valid_docx(self, simple_docx_bytes: bytes):
        # Should not raise
        validate_docx_magic_bytes(simple_docx_bytes)

    def test_too_short(self):
        with pytest.raises(InvalidDocxError, match="too short"):
            validate_docx_magic_bytes(b"PK")

    def test_invalid_header(self):
        with pytest.raises(InvalidDocxError, match="missing PK header"):
            validate_docx_magic_bytes(b"Not a DOCX file at all, just some text content here")

    def test_pdf_header_rejected(self):
        with pytest.raises(InvalidDocxError, match="missing PK header"):
            validate_docx_magic_bytes(b"%PDF-1.4 some content here")

    def test_generic_zip_rejected(self, generic_zip_bytes: bytes):
        """A valid ZIP that doesn't contain word/ entries should be rejected."""
        with pytest.raises(InvalidDocxError, match="does not contain word/ entries"):
            validate_docx_magic_bytes(generic_zip_bytes)

    def test_corrupted_zip(self):
        """ZIP header but corrupted content."""
        data = b"PK\x03\x04" + b"\x00" * 100
        with pytest.raises(InvalidDocxError, match="corrupt ZIP"):
            validate_docx_magic_bytes(data)


# =============================================================================
# Content-Type Validation Tests
# =============================================================================


class TestValidateDocxContentType:
    """Tests for validate_docx_content_type()."""

    def test_valid_docx_mime(self):
        validate_docx_content_type(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def test_valid_octet_stream(self):
        validate_docx_content_type("application/octet-stream")

    def test_valid_with_charset(self):
        validate_docx_content_type(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document; charset=utf-8"
        )

    def test_invalid_content_type(self):
        with pytest.raises(InvalidDocxError, match="Invalid Content-Type"):
            validate_docx_content_type("text/html")

    def test_pdf_content_type_rejected(self):
        with pytest.raises(InvalidDocxError, match="Invalid Content-Type"):
            validate_docx_content_type("application/pdf")

    def test_none_content_type(self):
        # Should not raise — falls through to magic byte validation
        validate_docx_content_type(None)


# =============================================================================
# DocxExtractor.extract() Tests
# =============================================================================


class TestDocxExtractorExtract:
    """Tests for DocxExtractor.extract()."""

    @pytest.mark.asyncio
    async def test_extract_simple_docx(self, simple_docx_bytes: bytes):
        extractor = DocxExtractor()
        result = await extractor.extract(simple_docx_bytes)

        assert result.success is True
        assert "Hello World" in result.text
        assert "test document" in result.text
        assert result.paragraph_count > 0

    @pytest.mark.asyncio
    async def test_extract_with_tables(self, docx_with_tables_bytes: bytes):
        extractor = DocxExtractor()
        result = await extractor.extract(docx_with_tables_bytes)

        assert result.success is True
        assert result.table_count == 1
        assert "Alice" in result.text
        assert "Bob" in result.text
        assert "New York" in result.text

    @pytest.mark.asyncio
    async def test_extract_empty_document(self, empty_docx_bytes: bytes):
        extractor = DocxExtractor()
        result = await extractor.extract(empty_docx_bytes)

        assert result.success is False
        assert result.has_warnings is True

    @pytest.mark.asyncio
    async def test_extract_from_bytesio(self, simple_docx_bytes: bytes):
        extractor = DocxExtractor()
        result = await extractor.extract(io.BytesIO(simple_docx_bytes))

        assert result.success is True
        assert "Hello World" in result.text

    @pytest.mark.asyncio
    async def test_extract_invalid_source_type(self):
        extractor = DocxExtractor()
        with pytest.raises(ValueError, match="must be bytes or BytesIO"):
            await extractor.extract("not bytes")  # type: ignore[arg-type]

    @pytest.mark.asyncio
    async def test_extract_invalid_magic_bytes(self):
        extractor = DocxExtractor()
        with pytest.raises(InvalidDocxError):
            await extractor.extract(b"Not a DOCX file, just some random text here")

    @pytest.mark.asyncio
    async def test_extract_skip_magic_validation(self, simple_docx_bytes: bytes):
        extractor = DocxExtractor()
        result = await extractor.extract(simple_docx_bytes, validate_magic=False)
        assert result.success is True

    @pytest.mark.asyncio
    async def test_size_limit_enforced(self, simple_docx_bytes: bytes):
        extractor = DocxExtractor(max_size=100)  # Very small limit
        with pytest.raises(DocxSizeError, match="exceeds limit"):
            await extractor.extract(simple_docx_bytes)

    @pytest.mark.asyncio
    async def test_size_limit_default(self):
        extractor = DocxExtractor()
        assert extractor.max_size == DEFAULT_MAX_DOCX_SIZE


# =============================================================================
# python-docx Not Installed Tests
# =============================================================================


class TestDocxExtractorGracefulDegradation:
    """Tests for graceful handling when python-docx is not installed."""

    @pytest.mark.asyncio
    async def test_extract_without_python_docx(self, simple_docx_bytes: bytes):
        """Extraction should raise RuntimeError when python-docx is missing."""
        with patch(
            "foundry_mcp.core.research.docx_extractor._get_docx_module",
            return_value=None,
        ):
            extractor = DocxExtractor()
            with pytest.raises(RuntimeError, match="python-docx is required"):
                await extractor.extract(simple_docx_bytes)


# =============================================================================
# SSRF Protection Tests
# =============================================================================


class TestDocxExtractorSSRF:
    """Tests for SSRF protection in extract_from_url()."""

    @pytest.mark.asyncio
    async def test_localhost_blocked(self):
        extractor = DocxExtractor()
        with pytest.raises(SSRFError, match="localhost"):
            await extractor.extract_from_url("http://localhost/doc.docx")

    @pytest.mark.asyncio
    async def test_private_ip_blocked(self):
        extractor = DocxExtractor()
        with pytest.raises(SSRFError, match="internal IP"):
            await extractor.extract_from_url("http://192.168.1.1/doc.docx")

    @pytest.mark.asyncio
    async def test_loopback_blocked(self):
        extractor = DocxExtractor()
        with pytest.raises(SSRFError, match="localhost"):
            await extractor.extract_from_url("http://127.0.0.1/doc.docx")

    @pytest.mark.asyncio
    async def test_metadata_endpoint_blocked(self):
        extractor = DocxExtractor()
        with pytest.raises(SSRFError, match="internal hostname"):
            await extractor.extract_from_url("http://169.254.169.254/latest/meta-data")

    @pytest.mark.asyncio
    async def test_invalid_scheme_blocked(self):
        extractor = DocxExtractor()
        with pytest.raises(SSRFError, match="Invalid URL scheme"):
            await extractor.extract_from_url("file:///etc/passwd")


# =============================================================================
# Thread Pool Execution Tests
# =============================================================================


class TestDocxExtractorThreadPool:
    """Tests verifying extraction runs in thread pool."""

    @pytest.mark.asyncio
    async def test_extract_uses_to_thread(self, simple_docx_bytes: bytes):
        """Verify that extraction dispatches to asyncio.to_thread."""
        import asyncio as _asyncio

        extractor = DocxExtractor()
        calls: list[tuple] = []

        original_to_thread = _asyncio.to_thread

        async def tracking_to_thread(func, /, *args, **kwargs):
            calls.append((func.__name__, args))
            return await original_to_thread(func, *args, **kwargs)

        with patch("foundry_mcp.core.research.docx_extractor.asyncio.to_thread", side_effect=tracking_to_thread):
            result = await extractor.extract(simple_docx_bytes)

        assert result.success is True
        assert len(calls) == 1
        assert calls[0][0] == "_extract_sync"


# =============================================================================
# Corrupted File Handling Tests
# =============================================================================


class TestDocxExtractorCorruptedFiles:
    """Tests for handling corrupted or malformed DOCX files."""

    @pytest.mark.asyncio
    async def test_corrupted_docx_content(self):
        """DOCX with valid magic bytes but corrupted internal content."""
        # Create a valid-looking ZIP with word/ entry but invalid DOCX content
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("word/document.xml", "<invalid>not real docx xml</invalid>")
            zf.writestr("[Content_Types].xml", "<Types></Types>")
        corrupted = buf.getvalue()

        extractor = DocxExtractor()
        # python-docx should fail to parse but we handle it gracefully
        result = await extractor.extract(corrupted)
        # Either succeeds with empty text or returns with warnings
        assert isinstance(result, DocxExtractionResult)
